"""
Generates 'Developer_Task_Breakdown.docx' — the detailed, hand-out-ready task
document for the Chatbot Platform, derived from context/task_queue.md and the
per-dev backlogs. Re-run after editing TASK DATA below to regenerate.

    python scripts/generate_task_doc.py
"""
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# TASK DATA  (edit here, then re-run)
# Each task: id, title, tag (S/P/D), size (S/M/L), depends, ai (Yes/Partial/NO),
#            refs, desc, notes (list), accept (list)
# ----------------------------------------------------------------------------

DEVS = [
    {
        "name": "Dev 1 — Architect + Senior",
        "scope": "DB schema, credential encryption, multi-tenancy, flow engine, "
                 "flow builder, auth scaffolding, handoff backend. Foundational "
                 "work — most Shared [S] items live here and unblock Dev 2 & Dev 3.",
        "epics": [
            ("Epic A — Project skeleton", [
                dict(id="D1-01", title="Init Django project + split settings",
                     tag="S", size="S", depends="—", ai="Yes", refs="architecture.md",
                     desc="Stand up the Django project with environment-driven, "
                          "split settings so dev and prod configs never mix.",
                     notes=["Create config/ project package with base/dev/prod settings modules.",
                            "Load secrets and DB/Redis URLs via django-environ from .env (never commit .env).",
                            "Provide .env.example with every key documented."],
                     accept=["manage.py runserver boots using .env values.",
                             "DEBUG, SECRET_KEY, DB and Redis come from environment, not hard-coded."]),
                dict(id="D1-02", title="Create 6 apps + register",
                     tag="S", size="S", depends="D1-01", ai="Yes", refs="architecture.md",
                     desc="Create the module boundaries the whole team builds inside.",
                     notes=["Apps: tenants, flows, conversations, channels, handoff, audit.",
                            "Add each to INSTALLED_APPS with an explicit AppConfig."],
                     accept=["All six apps import cleanly.",
                             "Project starts with the apps registered."]),
                dict(id="D1-03", title="Wire Celery + Redis",
                     tag="S", size="M", depends="D1-01", ai="Yes", refs="architecture.md C-02",
                     desc="Set up the async backbone that decouples webhook receipt "
                          "from flow execution.",
                     notes=["config/celery.py with app autodiscovery.",
                            "Redis as broker/result backend (from env).",
                            "Add a trivial ping task to prove the wiring."],
                     accept=["A Celery worker starts and runs the ping task end-to-end."]),
                dict(id="D1-04", title="Postgres connection",
                     tag="S", size="S", depends="D1-01", ai="Yes", refs="architecture.md",
                     desc="Point the project at PostgreSQL.",
                     notes=["DATABASES configured from env (psycopg)."],
                     accept=["manage.py migrate runs clean against a Postgres database."]),
            ]),
            ("Epic B — Data models + migrations (C-08)", [
                dict(id="D1-05", title="tenants model + admin register",
                     tag="S", size="M", depends="D1-02", ai="Partial", refs="architecture.md §Data models",
                     desc="The central tenant record that everything else hangs off.",
                     notes=["Fields: name, wa_phone_number, ig_account_id, handoff_enabled, "
                            "handoff_email, greeting_message, closing_message, is_active, created_at "
                            "(+ token fields added in D1-06).",
                            "Register in Django admin exposing config fields only (not raw tokens)."],
                     accept=["Product owner can create/edit a tenant's config in Django admin.",
                             "is_active toggles a tenant on/off."]),
                dict(id="D1-06", title="Token encryption at rest",
                     tag="S", size="L", depends="D1-05", ai="NO (SEC-02)", refs="CLAUDE.md §7, D-100",
                     desc="Store each tenant's Meta access tokens encrypted; the encryption "
                          "key lives outside the database. Resolves open decision D-100.",
                     notes=["Use Fernet (cryptography) or equivalent for wa_access_token / ig_access_token.",
                            "Key sourced from env / secrets manager — never stored in the DB (SEC-02).",
                            "Ensure tokens never appear in __repr__, admin, logs, or error output."],
                     accept=["Tokens are ciphertext at rest.",
                             "Key is not present anywhere in the database.",
                             "No code path logs a decrypted token.",
                             "decisions.md D-100 marked DECIDED with the chosen mechanism."]),
                dict(id="D1-07", title="flow_steps + flow_options models",
                     tag="S", size="M", depends="D1-02", ai="Partial", refs="architecture.md",
                     desc="The relational backbone of a conversation flow.",
                     notes=["flow_steps: tenant_id FK, label, message_text, is_start, is_terminal, created_at.",
                            "flow_options: step_id FK, button_label (max 20 chars), next_step_id FK nullable.",
                            "DB constraint: exactly one is_start=true per tenant; null next_step_id = terminal."],
                     accept=["button_label > 20 chars is rejected.",
                             "A tenant cannot have two start steps (enforced at DB level)."]),
                dict(id="D1-08", title="sessions + messages models",
                     tag="S", size="M", depends="D1-02", ai="Yes", refs="architecture.md",
                     desc="Per-conversation state and the message log.",
                     notes=["sessions: tenant_id, channel, customer_identifier, current_step_id, "
                            "status (active|completed|handed_off), started_at, updated_at.",
                            "messages: session_id, direction (inbound|outbound), content, channel, sent_at."],
                     accept=["A new conversation creates a new session (never reused).",
                             "Status enum restricted to the three allowed values."]),
                dict(id="D1-09", title="audit_logs + auto-record",
                     tag="S", size="M", depends="D1-05", ai="NO (SEC-05)", refs="CLAUDE.md §7, FR-19",
                     desc="Tamper-evident record of every product-owner action.",
                     notes=["Model: admin_user_id, action, entity_type, entity_id, diff (JSON), created_at.",
                            "Hook admin create/edit/delete to write an audit row automatically."],
                     accept=["Every config/flow change lands in audit_logs with actor + timestamp + diff."]),
                dict(id="D1-10", title="Migrations + constraints check",
                     tag="S", size="S", depends="D1-05..09", ai="Yes", refs="—",
                     desc="Lock the schema in with migrations and prove the key constraint.",
                     notes=["Generate and apply all migrations."],
                     accept=["All migrations apply cleanly.",
                             "A test proves the single-start-step constraint holds."]),
            ]),
            ("Epic C — Auth scaffolding (C-06)", [
                dict(id="D1-11", title="Admin auth (single PO account)",
                     tag="S", size="M", depends="D1-04", ai="NO (SEC-05)", refs="AP-01",
                     desc="Secure the admin surface for the single product-owner user.",
                     notes=["JWT or server session; no public registration.",
                            "All admin API endpoints reject unauthenticated requests."],
                     accept=["Login works for the product-owner account.",
                             "Unauthenticated API calls return 401."]),
                dict(id="D1-12", title="Example DRF endpoint (copy template)",
                     tag="S", size="M", depends="D1-11", ai="Yes", refs="CLAUDE.md §6",
                     desc="A reference vertical slice that Dev 3 copies for every CRUD endpoint.",
                     notes=["One resource with serializer + viewset + permission + test.",
                            "Document it as the canonical pattern in the repo."],
                     accept=["Endpoint passes its test and is documented as the template."]),
            ]),
            ("Epic D — Flow engine (C-03) — stateless, all state in DB", [
                dict(id="D1-13", title="Session resolver",
                     tag="P", size="M", depends="D1-08", ai="Partial", refs="FR-09, flow step 5",
                     desc="Find or create the session for an inbound event.",
                     notes=["Key on (tenant, channel, customer_identifier).",
                            "A new conversation starts a fresh session; no prior data referenced (D-05)."],
                     accept=["Returning customer restarts from greeting with a new session."]),
                dict(id="D1-14", title="Current-step renderer",
                     tag="P", size="M", depends="D1-07,D1-13", ai="Partial", refs="FR-05/06",
                     desc="Turn the session's current step into an outbound message payload.",
                     notes=["First inbound message sends the configured greeting before options.",
                            "Return step message_text plus its button options."],
                     accept=["Greeting precedes options on first contact.",
                             "Payload reflects the current step's options."]),
                dict(id="D1-15", title="Branching / advance",
                     tag="P", size="M", depends="D1-14", ai="NO (core logic)", refs="FR-07, flow step 7",
                     desc="Advance the session based on which button was tapped.",
                     notes=["Match the button reply to a flow_option of the current step.",
                            "Set current_step_id to the option's next_step_id."],
                     accept=["Tapping a button moves the session to the correct next step."]),
                dict(id="D1-16", title="Free-text handling",
                     tag="P", size="S", depends="D1-14", ai="Partial", refs="FR-08, flow step 6",
                     desc="Keep the flow on rails when a customer types free text.",
                     notes=["Resend the current step's message + buttons with a short prompt."],
                     accept=["Free text never advances or breaks the flow."]),
                dict(id="D1-17", title="Terminal handling",
                     tag="P", size="M", depends="D1-15", ai="NO (core logic)", refs="FR-10/12, flow step 8",
                     desc="Decide what happens when a path ends.",
                     notes=["Null next_step_id = terminal.",
                            "If handoff enabled → emit handoff offer; else send closing message and mark completed."],
                     accept=["Terminal with handoff on shows 'Talk to a human'.",
                             "Terminal with handoff off sends closing and marks session completed."]),
                dict(id="D1-18", title="Deleted-step recovery",
                     tag="P", size="S", depends="D1-15", ai="Yes", refs="ERR-04",
                     desc="Self-heal a session whose step was deleted under it.",
                     notes=["Detect current_step_id pointing to a missing step.",
                            "Reset to start step, resend greeting, log a recovery event."],
                     accept=["Orphaned session recovers to start and logs the event."]),
                dict(id="D1-19", title="Flow-engine unit tests",
                     tag="P", size="L", depends="D1-13..18", ai="Yes", refs="FR-05..10",
                     desc="Pin the engine's behaviour with tests across every branch.",
                     notes=["Cover greeting, branching, free-text, terminal, recovery."],
                     accept=["All flow-engine paths covered and green."]),
            ]),
            ("Epic E — Flow builder (C-07, FR-17) — hardest UI", [
                dict(id="D1-20", title="Decide live-flow-edit behaviour",
                     tag="D", size="M", depends="—", ai="NO", refs="D-103, AP-04",
                     desc="Decide and document how active sessions behave when a flow is "
                          "edited. Blocks the builder. Resolves D-103.",
                     notes=["Changes take effect on new sessions only (AP-04); define active-session handling."],
                     accept=["Behaviour documented in decisions.md; D-103 marked DECIDED."]),
                dict(id="D1-21", title="Step editor (form-based)",
                     tag="D", size="L", depends="D1-20,D1-07", ai="Partial", refs="AP-02",
                     desc="Form UI to create and edit steps and mark the start step.",
                     notes=["Fields: label + message_text; mark start; no drag-and-drop in v1."],
                     accept=["Product owner can create/edit steps and set the start step."]),
                dict(id="D1-22", title="Option editor",
                     tag="D", size="M", depends="D1-21", ai="Partial", refs="AP-02",
                     desc="Attach buttons to a step and wire them to next steps.",
                     notes=["Add/edit options: button_label + select next step from existing steps."],
                     accept=["Options can be added and pointed at any existing step."]),
                dict(id="D1-23", title="Flow validation gate",
                     tag="D", size="L", depends="D1-22", ai="NO (correctness)", refs="AP-03",
                     desc="Refuse to activate a structurally invalid flow.",
                     notes=["Validate: exactly one start, all next_step_id valid, no orphaned steps."],
                     accept=["Invalid flows cannot be activated; the reason is shown."]),
            ]),
            ("Epic F — Handoff backend (C-05)", [
                dict(id="D1-24", title="Handoff trigger service",
                     tag="D", size="L", depends="D1-17,D1-08", ai="Partial", refs="FR-11/13",
                     desc="Execute the human-handoff: notify staff, freeze the bot, "
                          "reassure the customer.",
                     notes=["Build payload: customer details, channel, timestamp, conversation path.",
                            "Send email (Dev 3 owns the template), mark session handed_off.",
                            "Send customer a confirmation, then stop bot responses in that session."],
                     accept=["Handoff email sent with full context.",
                             "Session marked handed_off and the bot goes silent.",
                             "Customer receives a confirmation message."]),
            ]),
        ],
    },
    {
        "name": "Dev 2 — Senior (Integrations)",
        "scope": "Webhook receiver, Meta signature validation, message queue, message "
                 "sender, WhatsApp 24-hour window, Instagram. Owns the Meta-facing edge "
                 "and the two open channel decisions (BSP, Instagram capability).",
        "epics": [
            ("Epic A — BSP & Meta setup", [
                dict(id="D2-01", title="Meta BSP selection",
                     tag="S", size="M", depends="—", ai="NO", refs="WA-01, D-101",
                     desc="Choose the WhatsApp Business Solution Provider before any WA "
                          "work starts; it dictates endpoint format and auth. Resolves D-101.",
                     notes=["Compare 360dialog / Twilio / Gupshup on pricing, API, onboarding time.",
                            "Note: Meta/BSP onboarding can take 1–2 weeks — start day one."],
                     accept=["BSP chosen with rationale; D-101 marked DECIDED.",
                             "Sandbox/number access requested."]),
                dict(id="D2-02", title="Webhook verification (GET challenge)",
                     tag="S", size="S", depends="D2-01", ai="Yes", refs="WA-04",
                     desc="Answer Meta's verification handshake during webhook setup.",
                     notes=["Respond to Meta's GET with the hub.challenge value and verify_token check."],
                     accept=["Meta successfully verifies the webhook endpoint."]),
            ]),
            ("Epic B — Webhook receiver (C-01)", [
                dict(id="D2-03", title="Inbound webhook endpoint",
                     tag="S", size="M", depends="D1-03,D2-02", ai="Partial", refs="C-01, ERR-01",
                     desc="Public HTTPS POST endpoint that always answers Meta fast and "
                          "offloads heavy work.",
                     notes=["Always return HTTP 200 to Meta (failures handled internally).",
                            "Respond within 5 seconds; all real processing is async."],
                     accept=["Endpoint returns 200 in well under 5s regardless of downstream state."]),
                dict(id="D2-04", title="HMAC-SHA256 signature validation",
                     tag="S", size="M", depends="D2-03", ai="NO (SEC-01)", refs="FR-03, SEC-01",
                     desc="Reject forged webhooks before any processing.",
                     notes=["Validate Meta's signature header with HMAC-SHA256 over the raw body.",
                            "Invalid → return 200, log, discard (no processing)."],
                     accept=["Tampered payloads are discarded and logged; valid ones pass."]),
                dict(id="D2-05", title="Payload extraction",
                     tag="S", size="M", depends="D2-04", ai="Yes", refs="flow step 2",
                     desc="Normalise channel, tenant, customer and message out of Meta's payload.",
                     notes=["Extract channel, tenant identifier, customer identifier, message payload.",
                            "Handle both WhatsApp and Instagram payload shapes."],
                     accept=["A normalised event object is produced for both channels."]),
                dict(id="D2-06", title="Tenant routing",
                     tag="S", size="M", depends="D2-05", ai="Partial", refs="FR-04, flow step 3",
                     desc="Route each inbound event to the correct tenant.",
                     notes=["Look up tenant by destination number / channel identifier.",
                            "Not found → log and discard, return 200 immediately."],
                     accept=["Events reach the right tenant; unknown destinations are safely dropped."]),
                dict(id="D2-07", title="Enqueue event",
                     tag="S", size="S", depends="D2-06", ai="Yes", refs="flow step 4",
                     desc="Hand the event to the queue and return.",
                     notes=["Push normalised event to Celery; return 200; all further work async."],
                     accept=["Event is queued and the request returns 200."]),
            ]),
            ("Epic C — Queue processing (C-02)", [
                dict(id="D2-08", title="Queue worker → flow engine",
                     tag="P", size="M", depends="D2-07,D1-19", ai="Partial", refs="flow steps 5–9",
                     desc="Consume queued events and drive the flow engine.",
                     notes=["Worker looks up/creates session and calls the flow engine.",
                            "Dispatch resulting outbound message to the sender."],
                     accept=["An inbound event produces the correct outbound response via the engine."]),
                dict(id="D2-09", title="Retry + dead-letter policy",
                     tag="P", size="M", depends="D2-08", ai="Partial", refs="ERR-02",
                     desc="Make processing resilient without dropping messages.",
                     notes=["Retry up to 3× with exponential backoff.",
                            "After 3 failures → dead-letter queue and alert the product owner."],
                     accept=["Transient failures retry; persistent ones dead-letter and alert."]),
            ]),
            ("Epic D — Message sender (C-04)", [
                dict(id="D2-10", title="Per-tenant sender abstraction",
                     tag="D", size="M", depends="D2-08,D1-06", ai="Partial", refs="C-04",
                     desc="Send outbound messages via the right Meta API with the right "
                          "tenant credentials.",
                     notes=["Resolve and decrypt per-tenant credentials at send time.",
                            "Channel-agnostic interface; channel-specific implementations behind it."],
                     accept=["Messages send using the correct tenant's credentials and channel."]),
                dict(id="D2-11", title="WhatsApp interactive messages",
                     tag="D", size="M", depends="D2-10", ai="Partial", refs="WA-02",
                     desc="Render WhatsApp reply buttons and list messages within Meta's limits.",
                     notes=["Reply buttons ≤ 3 options; list messages ≤ 10; labels ≤ 20 chars.",
                            "Auto-select list format when a step exceeds 3 options."],
                     accept=["Steps render as buttons or lists per option count, within limits."]),
                dict(id="D2-12", title="WhatsApp 24-hour window enforcement",
                     tag="D", size="M", depends="D2-11", ai="NO (constraint)", refs="WA-03",
                     desc="Respect WhatsApp's 24-hour interactive-messaging window.",
                     notes=["Check session.updated_at before sending.",
                            "If > 23h since last customer message → flag and do NOT send interactive.",
                            "v1: do not attempt templates outside the window — just flag and log."],
                     accept=["Outside-window interactive sends are blocked, flagged and logged."]),
                dict(id="D2-13", title="Outbound error handling",
                     tag="D", size="S", depends="D2-10", ai="Yes", refs="ERR-03",
                     desc="Handle Meta API send errors sanely.",
                     notes=["Log all send errors.",
                            "Do not retry 4xx; retry once on 5xx."],
                     accept=["4xx errors are not retried; 5xx retried once; all logged."]),
                dict(id="D2-14", title="Outbound message logging",
                     tag="D", size="S", depends="D2-10", ai="Yes", refs="flow step 9",
                     desc="Persist every outbound message.",
                     notes=["Write a messages row (direction=outbound) on each send."],
                     accept=["Every sent message is logged with channel and timestamp."]),
            ]),
            ("Epic E — Instagram (IG)", [
                dict(id="D2-15", title="Instagram capability assessment",
                     tag="S", size="M", depends="—", ai="NO", refs="IG-01, D-102",
                     desc="Assess Instagram DM API interactive-message support and report "
                          "to the product owner for approval. Resolves D-102.",
                     notes=["Determine native button/quick-reply support.",
                            "If unsupported, propose numbered-text-menu fallback for PO approval."],
                     accept=["Findings + recommended UX approved by PO; D-102 marked DECIDED."]),
                dict(id="D2-16", title="Instagram permissions / OAuth",
                     tag="D", size="M", depends="D2-15", ai="Partial", refs="IG-03",
                     desc="Connect each tenant's Instagram account with the needed scopes.",
                     notes=["Grant instagram_manage_messages and pages_messaging via Meta OAuth.",
                            "Requires a linked Facebook Page per tenant."],
                     accept=["A tenant's IG account can send/receive messages through the platform."]),
                dict(id="D2-17", title="Instagram renderer / fallback",
                     tag="D", size="L", depends="D2-16,D2-12", ai="Partial", refs="IG-02",
                     desc="Render flows on Instagram, using the approved fallback if buttons "
                          "are unavailable.",
                     notes=["If no buttons: parse '1'/'2'/'3' as option selections — Instagram only.",
                            "Must not change WhatsApp logic."],
                     accept=["IG sessions navigate the flow via the approved UX.",
                             "WhatsApp behaviour is unaffected."]),
                dict(id="D2-18", title="Channel integration tests",
                     tag="D", size="L", depends="D2-17,D2-14", ai="Yes", refs="FR-01..04",
                     desc="Prove inbound→outbound works on both channels.",
                     notes=["Simulated webhook → engine → sender for WhatsApp and Instagram."],
                     accept=["End-to-end channel tests pass for both WhatsApp and Instagram."]),
            ]),
        ],
    },
    {
        "name": "Dev 3 — Junior (Admin & Config)",
        "scope": "Admin CRUD, per-client config screens, conversation-log viewer, handoff "
                 "email content, seed/test data, onboarding runbook. Low blast-radius, "
                 "screen-verifiable work; security and flow-state logic stay with the Seniors.",
        "epics": [
            ("Epic A — Admin CRUD & config", [
                dict(id="D3-01", title="Tenant create / edit / deactivate",
                     tag="P", size="M", depends="D1-05,D1-12", ai="Yes", refs="FR-15, AP-01",
                     desc="Let the product owner manage business-client accounts.",
                     notes=["Fields: name, WA number, Instagram account ID, handoff email, active status.",
                            "Build on Django admin / the D1-12 endpoint template."],
                     accept=["PO can create, edit and deactivate a client.",
                             "Deactivated clients are clearly marked."]),
                dict(id="D3-02", title="Per-client config screens",
                     tag="P", size="M", depends="D3-01", ai="Yes", refs="FR-16",
                     desc="Expose the per-client conversational settings.",
                     notes=["Edit greeting message, closing message, handoff toggle, handoff email."],
                     accept=["Each setting saves and is reflected by the bot's behaviour."]),
                dict(id="D3-03", title="Config validation",
                     tag="P", size="S", depends="D3-02", ai="Yes", refs="FR-16",
                     desc="Prevent obviously broken configs.",
                     notes=["Require handoff_email when handoff is enabled; validate email format."],
                     accept=["Enabling handoff without a valid email is blocked with a clear message."]),
            ]),
            ("Epic B — Conversation logs", [
                dict(id="D3-04", title="Conversation-log viewer",
                     tag="P", size="M", depends="D1-08", ai="Yes", refs="FR-18",
                     desc="Read-only view of message history per client.",
                     notes=["List messages with direction, channel and timestamps, filterable by tenant.",
                            "Read-only — no editing."],
                     accept=["PO can browse a client's message history with timestamps and channel."]),
                dict(id="D3-05", title="Path-taken display",
                     tag="P", size="M", depends="D3-04", ai="Partial", refs="FR-18",
                     desc="Show the sequence of steps a session walked.",
                     notes=["Reconstruct and display the ordered path for a session."],
                     accept=["A session's button path is visible end-to-end."]),
                dict(id="D3-06", title="Audit-log viewer",
                     tag="P", size="S", depends="D1-09", ai="Yes", refs="FR-19",
                     desc="Surface the audit trail for review.",
                     notes=["Read-only list of audit_logs with actor, action, entity, timestamp."],
                     accept=["Admin actions are reviewable in a read-only audit view."]),
            ]),
            ("Epic C — Handoff content", [
                dict(id="D3-07", title="Handoff email template",
                     tag="D", size="M", depends="D1-24", ai="Yes", refs="FR-11",
                     desc="Author the email the human agent receives on handoff.",
                     notes=["Include customer details, channel, timestamp, conversation path.",
                            "Plain, scannable layout; no customer PII beyond what's needed."],
                     accept=["Handoff email is clear, complete and renders in common clients."]),
                dict(id="D3-08", title="Customer confirmation message",
                     tag="D", size="S", depends="D1-24", ai="Yes", refs="FR-13",
                     desc="Write the message the customer sees after requesting a human.",
                     notes=["Reassure them a human will follow up; set expectations."],
                     accept=["Confirmation copy approved and wired into the handoff flow."]),
            ]),
            ("Epic D — Fixtures & enablement", [
                dict(id="D3-09", title="Demo seed data",
                     tag="P", size="M", depends="D1-10", ai="Yes", refs="architecture.md",
                     desc="A realistic demo tenant to develop and demo against.",
                     notes=["One tenant + a sample flow: greeting → two branches → terminal."],
                     accept=["A single command seeds a working demo tenant and flow."]),
                dict(id="D3-10", title="Test fixtures / factories",
                     tag="P", size="M", depends="D1-10", ai="Yes", refs="—",
                     desc="Reusable factories so every dev's tests have data.",
                     notes=["factory-boy factories for tenants, flows, steps, options, sessions."],
                     accept=["Other devs can build test objects from shared factories."]),
                dict(id="D3-11", title="Client onboarding runbook",
                     tag="P", size="S", depends="D3-02", ai="Yes", refs="FR-15/16",
                     desc="A short guide the product owner follows to onboard a client.",
                     notes=["Step-by-step: create tenant, enter credentials, set messages, build flow, go live."],
                     accept=["A non-engineer can onboard a client by following the runbook."]),
            ]),
        ],
    },
]

# ----------------------------------------------------------------------------
# RENDERING
# ----------------------------------------------------------------------------

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
TAG_NAMES = {"S": "[S] Shared", "P": "[P] Parallel", "D": "[D] Dependent"}
SIZE_NAMES = {"S": "S (~½ day)", "M": "M (~1 day)", "L": "L (~2 days)"}


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def meta_table(doc, t):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Tag", TAG_NAMES[t["tag"]]),
        ("Size", SIZE_NAMES[t["size"]]),
        ("Depends on", t["depends"]),
        ("AI-junior?", t["ai"]),
        ("Spec refs", t["refs"]),
    ]
    first = table.rows[0]
    first.cells[0].text, first.cells[1].text = rows[0]
    for k, v in rows[1:]:
        c = table.add_row().cells
        c[0].text, c[1].text = k, v
    for r in table.rows:
        shade(r.cells[0], "EAEEF3")
        for p in r.cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
        r.cells[0].width = Inches(1.4)
        r.cells[1].width = Inches(4.8)
    doc.add_paragraph()


def add_task(doc, t):
    h = doc.add_heading(level=3)
    run = h.add_run(f"{t['id']} · {t['title']}")
    run.font.color.rgb = NAVY
    meta_table(doc, t)

    p = doc.add_paragraph()
    p.add_run("Description. ").bold = True
    p.add_run(t["desc"])

    doc.add_paragraph("Technical notes", style="Intense Quote")
    for n in t["notes"]:
        doc.add_paragraph(n, style="List Bullet")

    doc.add_paragraph("Acceptance criteria", style="Intense Quote")
    for a in t["accept"]:
        doc.add_paragraph(f"☐  {a}", style="List Bullet")
    doc.add_paragraph()


def add_mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    return p


def simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade(hdr[i], "1F3A5F")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def add_dependency_section(doc):
    doc.add_page_break()
    h = doc.add_heading("Execution order & dependencies", level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
    doc.add_paragraph(
        "The 53 tasks are neither all independent nor a single chain. They form three "
        "layers: a foundational series that gates everything, three parallel streams that "
        "run concurrently, and a dependent finish that converges. Build [S] first, run [P] "
        "in parallel, then complete [D]."
    )

    doc.add_heading("Layer 1 — Foundation [S] (Dev 1, first)", level=2)
    doc.add_paragraph(
        "Largely sequential. Blocks Dev 2 and Dev 3 until done — ship D1-05 + D1-12 "
        "early to unblock the others fast."
    )
    add_mono(doc, "D1-01 init → D1-02 apps → D1-05 tenants → D1-07/08 models → D1-10 migrations")
    add_mono(doc, "D1-05 → D1-06 encryption        D1-04 → D1-11 auth → D1-12 endpoint template")

    doc.add_heading("Layer 2 — Three parallel streams [P] (run concurrently)", level=2)
    doc.add_paragraph("Independent of each other; each is an internal series.")
    add_mono(doc, "Dev 1 engine : D1-13 → D1-14 → D1-15 → D1-17 → D1-19")
    add_mono(doc, "Dev 2 webhook: D2-03 → D2-04 → D2-05 → D2-06 → D2-07 → D2-08 → D2-09")
    add_mono(doc, "Dev 3 admin  : D3-01 → D3-02 → D3-03   and   D3-04 → D3-05")

    doc.add_heading("Layer 3 — Dependent finish [D] (converges)", level=2)
    add_mono(doc, "message sender : needs engine D1-19 + queue D2-08 + encryption D1-06")
    add_mono(doc, "flow builder   : needs decision D-103/D1-20 + models D1-07")
    add_mono(doc, "handoff        : needs engine-terminal D1-17 + email template D3-07")
    add_mono(doc, "Instagram      : needs WA sender D2-12 + IG decision D2-15")

    doc.add_heading("Critical path (longest chain)", level=2)
    doc.add_paragraph(
        "Runs entirely through Dev 1's foundation, then the engine, then the sender, "
        "then Instagram — roughly 16 working days of strictly sequential work. Any slip "
        "here slips the whole project; Dev 1 is the bottleneck."
    )
    add_mono(doc, "D1-01 → D1-02 → D1-07/08 → D1-13 → D1-14 → D1-15 → D1-17 → D1-19")
    add_mono(doc, "      → D2-08 → D2-10 → D2-11 → D2-12 → D2-17 → D2-18")

    doc.add_heading("Cross-developer handoffs (the scheduling risk)", level=2)
    doc.add_paragraph("Where one person waits on another. Everything traces back to Dev 1.")
    simple_table(
        doc,
        ["Blocked task", "Waits for", "Direction"],
        [
            ["D2-08 queue worker", "D1-19 flow engine tested", "Dev 2 ← Dev 1"],
            ["D2-10 message sender", "D1-06 token encryption", "Dev 2 ← Dev 1"],
            ["D3-01 admin CRUD", "D1-05 model + D1-12 template", "Dev 3 ← Dev 1"],
            ["D3-07 handoff email", "D1-24 handoff backend", "Dev 3 ← Dev 1"],
            ["D3-09 seed data", "D1-10 migrations", "Dev 3 ← Dev 1"],
            ["D2-17 Instagram", "D2-12 WA sender + D2-15 IG decision", "Dev 2 (internal)"],
            ["D1-21 flow builder", "D1-20 live-flow decision", "Dev 1 (internal)"],
        ],
        widths=[2.3, 2.9, 1.6],
    )
    doc.add_paragraph(
        "Note: Dev 3 (Junior) is blocked earliest — nearly all their work needs Dev 1's "
        "models + endpoint template. In week 1, pair Dev 3 on the foundation or front-load "
        "non-code tasks (D3-11 runbook, UI mockups)."
    )

    doc.add_heading("Open-decision blockers (resolve in Week 0)", level=2)
    doc.add_paragraph("Each sits at the head of a chain; the chain stalls until it is decided.")
    simple_table(
        doc,
        ["Decision", "Blocks", "Owner / when"],
        [
            ["D-100 credential encryption", "D1-06 → message sender", "Dev 1 · Week 0"],
            ["D-101 Meta BSP selection", "D2-01 → all WhatsApp work", "PO + Dev 2 · Week 0"],
            ["D-102 Instagram capabilities", "D2-15 → Instagram renderer", "Dev 2 + PO · Week 0"],
            ["D-103 live-flow-edit behaviour", "D1-20 → flow builder", "Dev 1 · before builder"],
        ],
        widths=[2.6, 2.6, 1.6],
    )

    doc.add_heading("Suggested execution phases", level=2)
    simple_table(
        doc,
        ["Phase", "Focus", "Tasks"],
        [
            ["0 — Week 0", "Resolve decisions + scaffold", "D-100..103, D1-01..04, D2-01, D2-15"],
            ["1 — Foundation", "Models, encryption, auth", "D1-05..12"],
            ["2 — Parallel", "Engine + webhook + admin", "D1-13..19, D2-03..09, D3-01..06"],
            ["3 — Convergent", "Sender, builder, handoff, IG", "D2-10..18, D1-20..24, D3-07..11"],
        ],
        widths=[1.5, 2.5, 2.8],
    )


def build():
    doc = Document()

    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Chatbot Platform")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Developer Task Breakdown")
    rs.font.size = Pt(16)
    rs.font.color.rgb = GREY
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"WhatsApp + Instagram  ·  Multi-tenant  ·  Rule-based\n"
        f"Version 1.0  ·  Generated {date.today().isoformat()}  ·  Stack: Django + DRF\n"
        f"Team: 3 developers (2 Senior + 1 Junior) + AI juniors"
    ).font.color.rgb = GREY
    doc.add_paragraph()

    # --- How to use ---
    doc.add_heading("How to use this document", level=1)
    doc.add_paragraph(
        "These tasks follow the PIER workflow (Plan → Identify → Execute → Review) "
        "from CLAUDE.md. Every task runs full PIER; effort is proportionate to size. "
        "Tasks are grouped by developer, then by epic. Build [S] Shared work first, "
        "run [P] Parallel work concurrently, then complete [D] Dependent work."
    )
    for label, text in [
        ("Tags", "[S] Shared = foundational, build first · [P] Parallel = independent · "
                 "[D] Dependent = needs prior tasks."),
        ("Size", "S ≈ ½ day · M ≈ 1 day · L ≈ 2 days."),
        ("AI-junior?", "Yes = safe to delegate boilerplate to an AI sub-agent (human reviews) · "
                       "Partial = scaffold with AI, human owns logic · "
                       "NO = security or core logic; humans only, never delegated."),
        ("Golden Rule", "AI never writes code until a human says “IMPLEMENT.” "
                        "Security guardrails (CLAUDE.md §7) are never delegated."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)

    # --- Execution order & dependencies ---
    add_dependency_section(doc)

    # --- Per developer ---
    for dev in DEVS:
        doc.add_page_break()
        h = doc.add_heading(dev["name"], level=1)
        for run in h.runs:
            run.font.color.rgb = NAVY
        sc = doc.add_paragraph()
        sc.add_run("Scope. ").bold = True
        sc.add_run(dev["scope"])

        total = sum(len(tasks) for _, tasks in dev["epics"])
        doc.add_paragraph(f"{total} tasks across {len(dev['epics'])} epics.").italic = True

        for epic_name, tasks in dev["epics"]:
            eh = doc.add_heading(epic_name, level=2)
            for run in eh.runs:
                run.font.color.rgb = GREY
            for t in tasks:
                add_task(doc, t)

    out = "Developer_Task_Breakdown.docx"
    doc.save(out)
    grand = sum(len(tasks) for dev in DEVS for _, tasks in dev["epics"])
    print(f"Wrote {out} — {grand} tasks across {len(DEVS)} developers.")


if __name__ == "__main__":
    build()
